from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file, flash
from flask_cors import CORS
from markupsafe import Markup
import sqlite3
import os
import random
import string
from datetime import datetime
import qrcode
import io
import base64
import shutil
import csv
from docx import Document
from docx.shared import Inches, Pt, Cm
from io import BytesIO
import json
import re
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import zipfile
from pathlib import Path
import spotipy
from spotipy.oauth2 import SpotifyOAuth

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from PIL import Image as PILImage
import requests
from meross_iot.http_api import MerossHttpClient
from meross_iot.manager import MerossManager
from meross_iot.controller.mixins.toggle import ToggleXMixin
from meross_iot.model.enums import OnlineStatus
from concurrent.futures import ThreadPoolExecutor
import threading

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}}, supports_credentials=True)
app.secret_key = 'your-secret-key'

VERSION = "5.0"
LATEST_VERSION = None
UPDATE_AVAILABLE = False

#------------------------Spotify Part--------------------------------------

load_dotenv()

MEROSS_EMAIL = os.getenv("MEROSS_EMAIL")
MEROSS_PASSWORD = os.getenv("MEROSS_PASSWORD")
API_BASE = "https://iotx-eu.meross.com"



SPOTIPY_CLIENT_ID = os.getenv("SPOTIPY_CLIENT_ID")
SPOTIPY_CLIENT_SECRET = os.getenv("SPOTIPY_CLIENT_SECRET")
SPOTIPY_REDIRECT_URI = os.getenv("SPOTIPY_REDIRECT_URI")

SCOPE = 'user-read-currently-playing user-read-playback-state user-modify-playback-state'

@app.route('/clear-spotify-cache')
def clear_spotify_cache():
    session.clear()
    return "Cache gelöscht! Gehe zu <a href='/spotify'>/spotify</a>"

def get_spotify_client():
    """Erstellt einen authentifizierten Spotify Client"""
    cache_handler = spotipy.cache_handler.FlaskSessionCacheHandler(session)
    auth_manager = SpotifyOAuth(
        client_id=SPOTIPY_CLIENT_ID,
        client_secret=SPOTIPY_CLIENT_SECRET,
        redirect_uri=SPOTIPY_REDIRECT_URI,
        scope=SCOPE,
        cache_handler=cache_handler,
        show_dialog=True
    )
    
    if not auth_manager.validate_token(cache_handler.get_cached_token()):
        return None
    
    return spotipy.Spotify(auth_manager=auth_manager)

@app.route('/spotify')
def spotify_auth():
    """Spotify Authentifizierung"""
    sp = get_spotify_client()
    if sp is None:
        cache_handler = spotipy.cache_handler.FlaskSessionCacheHandler(session)
        auth_manager = SpotifyOAuth(
            client_id=SPOTIPY_CLIENT_ID,
            client_secret=SPOTIPY_CLIENT_SECRET,
            redirect_uri=SPOTIPY_REDIRECT_URI,
            scope=SCOPE,
            cache_handler=cache_handler,
            show_dialog=True
        )
        auth_url = auth_manager.get_authorize_url()
        return redirect(auth_url)
    
    return redirect('/player')

@app.route('/callback')
def callback():
    """Callback nach Spotify Authentifizierung"""
    cache_handler = spotipy.cache_handler.FlaskSessionCacheHandler(session)
    auth_manager = SpotifyOAuth(
        client_id=SPOTIPY_CLIENT_ID,
        client_secret=SPOTIPY_CLIENT_SECRET,
        redirect_uri=SPOTIPY_REDIRECT_URI,
        scope=SCOPE,
        cache_handler=cache_handler,
        show_dialog=True
    )
    
    if request.args.get('code'):
        auth_manager.get_access_token(request.args.get('code'))
        return redirect('/player')
    
    return redirect('/')

@app.route('/player')
def player():
    """Player Seite"""
    sp = get_spotify_client()
    if sp is None:
        return redirect('/spotify')
    
    return render_template('player.html')

@app.route('/api/current')
def current_track():
    """API Endpoint für aktuell gespieltes Lied"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        current = sp.current_playback()
        if current and current['item']:
            track = current['item']
            return jsonify({
                'is_playing': current['is_playing'],
                'track_name': track['name'],
                'artist_name': ', '.join([artist['name'] for artist in track['artists']]),
                'album_name': track['album']['name'],
                'album_image': track['album']['images'][0]['url'] if track['album']['images'] else None,
                'duration_ms': track['duration_ms'],
                'progress_ms': current['progress_ms'],
                'volume': current['device']['volume_percent'] if current['device'] else 50
            })
        else:
            return jsonify({'is_playing': False, 'message': 'Kein Track wird gerade abgespielt'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/next', methods=['POST'])
def next_track():
    """Nächster Track"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        sp.next_track()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/previous', methods=['POST'])
def previous_track():
    """Vorheriger Track"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        sp.previous_track()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/play', methods=['POST'])
def play():
    """Wiedergabe starten"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        sp.start_playback()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pause', methods=['POST'])
def pause():
    """Wiedergabe pausieren"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        sp.pause_playback()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/volume', methods=['POST'])
def set_volume():
    """Lautstärke ändern"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        volume = int(request.json.get('volume', 50))
        volume = max(0, min(100, volume))
        sp.volume(volume)
        return jsonify({'success': True, 'volume': volume})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/search')
def search_tracks():
    """Tracks suchen"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    query = request.args.get('q', '')
    if not query or len(query) < 2:
        return jsonify({'tracks': []})
    
    try:
        results = sp.search(q=query, type='track', limit=10)
        tracks = []
        for item in results['tracks']['items']:
            tracks.append({
                'id': item['id'],
                'uri': item['uri'],
                'name': item['name'],
                'artist': ', '.join([artist['name'] for artist in item['artists']]),
                'album': item['album']['name'],
                'image': item['album']['images'][0]['url'] if item['album']['images'] else None,
                'duration_ms': item['duration_ms']
            })
        return jsonify({'tracks': tracks})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/play-track', methods=['POST'])
def play_track():
    """Bestimmten Track abspielen"""
    sp = get_spotify_client()
    if sp is None:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        track_uri = request.json.get('uri')
        if track_uri:
            sp.start_playback(uris=[track_uri])
            return jsonify({'success': True})
        return jsonify({'error': 'No URI provided'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/spotifylogout')
def spotify_logout():
    """Logout"""
    session.clear()
    return redirect('/spotify')
#-----------------------Spotify end ---------------------------------------

def check_version():
    global LATEST_VERSION, UPDATE_AVAILABLE
    try:
        response = requests.get("https://raw.githubusercontent.com/Matti-Krebelder/DMS/refs/heads/main/version.txt", timeout=5)
        if response.status_code == 200:
            latest_version = response.text.strip()
            LATEST_VERSION = latest_version
            if latest_version != VERSION:
                UPDATE_AVAILABLE = True
                print(f"Version {VERSION} ist veraltet. Aktuelle Version: {latest_version}")
            else:
                UPDATE_AVAILABLE = False
                print(f"Version {VERSION} ist aktuell.")
        else:
            print("Fehler beim Abrufen der Version.")
    except Exception as e:
        print(f"Fehler beim Überprüfen der Version: {e}")

def init_user_db():
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id TEXT PRIMARY KEY, name TEXT NOT NULL)''')
    c.execute('''CREATE TABLE IF NOT EXISTS lager
                 (id TEXT PRIMARY KEY, name TEXT NOT NULL, created_by TEXT,
                   access_users TEXT, system_type DEFAULT 'personal')''')
    c.execute("INSERT OR IGNORE INTO users VALUES ('CKS.EXampleid', 'Matti')")
    c.execute("INSERT OR IGNORE INTO users VALUES ('CKS-Example', 'Hubert')")
    c.execute("INSERT OR IGNORE INTO users VALUES ('CKS-Exampledsa', 'Admin')")
    c.execute("INSERT OR IGNORE INTO users VALUES ('CKS-7sdfuh-dfi', 'Christoffer Rentsch')")
    c.execute("INSERT OR IGNORE INTO users VALUES ('CKS-udzsfzewliuhd', 'Steffen Mascher')")
    conn.commit()
    conn.close()

def create_warehouse_db(lager_id):
    conn = sqlite3.connect(f'{lager_id}.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE geraete
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                   name TEXT NOT NULL,
                   barcode TEXT UNIQUE NOT NULL,
                   lagerplatz TEXT NOT NULL,
                   status TEXT DEFAULT 'verfügbar',
                   beschreibung TEXT,
                   seriennummer TEXT,
                   modell TEXT,
                   instrumentenart TEXT,
                   inventarnummer TEXT,
                   kaufdatum TEXT,
                   preis REAL,
                   quantity INTEGER DEFAULT 1,
                   hersteller TEXT)''')
    c.execute('''CREATE TABLE ausleihen
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                   ausleih_id TEXT NOT NULL,
                   mitarbeiter_id TEXT NOT NULL,
                   mitarbeiter_name TEXT NOT NULL,
                   zielort TEXT NOT NULL,
                   datum TEXT NOT NULL,
                   rueckgabe_qr TEXT NOT NULL,
                   status TEXT DEFAULT 'ausgeliehen',
                   email TEXT,
                   klasse TEXT)''')
    c.execute('''CREATE TABLE ausleih_details
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                   ausleih_id TEXT NOT NULL,
                   geraet_id INTEGER NOT NULL,
                   geraet_barcode TEXT NOT NULL,
                   quantity INTEGER DEFAULT 1,
                   FOREIGN KEY(ausleih_id) REFERENCES ausleihen(ausleih_id),
                   FOREIGN KEY(geraet_id) REFERENCES geraete(id))''')
    c.execute('''CREATE TABLE label_layouts
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                   name TEXT NOT NULL,
                   layout_data TEXT NOT NULL,
                   is_default INTEGER DEFAULT 0,
                   created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                   updated_at TEXT DEFAULT CURRENT_TIMESTAMP)''')
    conn.commit()
    conn.close()

import sqlite3
import shutil
from datetime import datetime
import os

def backup_database(db_path, operation="auto_migration"):
    """Create a backup of the database before making changes"""
    os.makedirs('backups', exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_path = f'backups/{timestamp}_{operation}_{os.path.basename(db_path)}'
    shutil.copy(db_path, backup_path)
    print(f"Backup erstellt: {backup_path}")
    return backup_path

def get_table_columns(conn, table_name):
    """Get all columns of a table"""
    cursor = conn.cursor()
    cursor.execute(f"PRAGMA table_info({table_name})")
    return {row[1]: row[2] for row in cursor.fetchall()}  # {column_name: column_type}

def add_missing_column(conn, table_name, column_name, column_type, default_value=None):
    """Add a missing column to a table"""
    cursor = conn.cursor()
    try:
        sql = f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_type}"
        if default_value is not None:
            sql += f" DEFAULT {default_value}"
        cursor.execute(sql)
        conn.commit()
        print(f"Spalte hinzugefügt: {table_name}.{column_name} ({column_type})")
        return True
    except sqlite3.Error as e:
        print(f"Fehler beim Hinzufügen der Spalte {column_name}: {e}")
        return False

def check_and_migrate_warehouse_db(lager_id):
    """Check and migrate a warehouse database with all required columns"""
    db_path = f'{lager_id}.db'
    
    if not os.path.exists(db_path):
        print(f"Datenbank {db_path} existiert nicht")
        return False
    
    # Backup vor Migration
    backup_database(db_path, "migration")
    
    conn = sqlite3.connect(db_path)
    
    try:
        # Definiere erwartete Spalten für jede Tabelle
        expected_schema = {
            'geraete': {
                'id': 'INTEGER PRIMARY KEY AUTOINCREMENT',
                'name': 'TEXT NOT NULL',
                'barcode': 'TEXT UNIQUE NOT NULL',
                'lagerplatz': 'TEXT NOT NULL',
                'status': 'TEXT DEFAULT "verfügbar"',
                'beschreibung': 'TEXT',
                'seriennummer': 'TEXT',
                'modell': 'TEXT',
                'instrumentenart': 'TEXT',
                'inventarnummer': 'TEXT',
                'kaufdatum': 'TEXT',
                'preis': 'REAL',
                'quantity': 'INTEGER DEFAULT 1',
                'hersteller': 'TEXT'
            },
            'ausleihen': {
                'id': 'INTEGER PRIMARY KEY AUTOINCREMENT',
                'ausleih_id': 'TEXT NOT NULL',
                'mitarbeiter_id': 'TEXT NOT NULL',
                'mitarbeiter_name': 'TEXT NOT NULL',
                'zielort': 'TEXT NOT NULL',
                'datum': 'TEXT NOT NULL',
                'rueckgabe_qr': 'TEXT NOT NULL',
                'status': 'TEXT DEFAULT "ausgeliehen"',
                'email': 'TEXT',
                'klasse': 'TEXT'
            },
            'ausleih_details': {
                'id': 'INTEGER PRIMARY KEY AUTOINCREMENT',
                'ausleih_id': 'TEXT NOT NULL',
                'geraet_id': 'INTEGER NOT NULL',
                'geraet_barcode': 'TEXT NOT NULL',
                'quantity': 'INTEGER DEFAULT 1'
            },
            'label_layouts': {
                'id': 'INTEGER PRIMARY KEY AUTOINCREMENT',
                'name': 'TEXT NOT NULL',
                'layout_data': 'TEXT NOT NULL',
                'is_default': 'INTEGER DEFAULT 0',
                'created_at': 'TEXT DEFAULT CURRENT_TIMESTAMP',
                'updated_at': 'TEXT DEFAULT CURRENT_TIMESTAMP'
            }
        }
        
        changes_made = False
        
        # Prüfe jede Tabelle
        for table_name, expected_columns in expected_schema.items():
            # Prüfe ob Tabelle existiert
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
            
            if not cursor.fetchone():
                print(f"Tabelle {table_name} existiert nicht - wird erstellt")
                # Tabelle erstellen (hier könntest du die CREATE TABLE Statements einfügen)
                if table_name == 'geraete':
                    cursor.execute('''CREATE TABLE geraete
                                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                       name TEXT NOT NULL,
                                       barcode TEXT UNIQUE NOT NULL,
                                       lagerplatz TEXT NOT NULL,
                                       status TEXT DEFAULT 'verfügbar',
                                       beschreibung TEXT,
                                       seriennummer TEXT,
                                       modell TEXT,
                                       instrumentenart TEXT,
                                       inventarnummer TEXT,
                                       kaufdatum TEXT,
                                       preis REAL,
                                       quantity INTEGER DEFAULT 1,
                                       hersteller TEXT)''')
                elif table_name == 'ausleihen':
                    cursor.execute('''CREATE TABLE ausleihen
                                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                       ausleih_id TEXT NOT NULL,
                                       mitarbeiter_id TEXT NOT NULL,
                                       mitarbeiter_name TEXT NOT NULL,
                                       zielort TEXT NOT NULL,
                                       datum TEXT NOT NULL,
                                       rueckgabe_qr TEXT NOT NULL,
                                       status TEXT DEFAULT 'ausgeliehen',
                                       email TEXT,
                                       klasse TEXT)''')
                elif table_name == 'ausleih_details':
                    cursor.execute('''CREATE TABLE ausleih_details
                                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                       ausleih_id TEXT NOT NULL,
                                       geraet_id INTEGER NOT NULL,
                                       geraet_barcode TEXT NOT NULL,
                                       quantity INTEGER DEFAULT 1,
                                       FOREIGN KEY(ausleih_id) REFERENCES ausleihen(ausleih_id),
                                       FOREIGN KEY(geraet_id) REFERENCES geraete(id))''')
                elif table_name == 'label_layouts':
                    cursor.execute('''CREATE TABLE label_layouts
                                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                       name TEXT NOT NULL,
                                       layout_data TEXT NOT NULL,
                                       is_default INTEGER DEFAULT 0,
                                       created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                                       updated_at TEXT DEFAULT CURRENT_TIMESTAMP)''')
                conn.commit()
                changes_made = True
                continue
            
            # Hole aktuelle Spalten
            current_columns = get_table_columns(conn, table_name)
            
            # Prüfe fehlende Spalten
            for column_name, column_type in expected_columns.items():
                if column_name not in current_columns:
                    print(f"Fehlende Spalte erkannt: {table_name}.{column_name}")
                    
                    # Extrahiere DEFAULT-Wert wenn vorhanden
                    default_value = None
                    if 'DEFAULT' in column_type:
                        parts = column_type.split('DEFAULT')
                        column_type = parts[0].strip()
                        default_value = parts[1].strip().strip("'\"")
                        if default_value.upper() != 'CURRENT_TIMESTAMP':
                            default_value = f"'{default_value}'"
                    
                    if add_missing_column(conn, table_name, column_name, column_type, default_value):
                        changes_made = True
        
        if changes_made:
            print(f"Migration für {lager_id} abgeschlossen")
        else:
            print(f"Keine Änderungen für {lager_id} erforderlich")
        
        return True
        
    except Exception as e:
        print(f"Fehler bei der Migration: {e}")
        return False
    finally:
        conn.close()

def check_and_migrate_users_db():
    """Check and migrate the users database"""
    db_path = 'users.db'
    
    if not os.path.exists(db_path):
        print("users.db existiert nicht - wird initialisiert")
        init_user_db()
        return
    
    backup_database(db_path, "migration")
    
    conn = sqlite3.connect(db_path)
    
    try:
        # Erwartete Spalten für users.db
        expected_schema = {
            'users': {
                'id': 'TEXT PRIMARY KEY',
                'name': 'TEXT NOT NULL'
            },
            'lager': {
                'id': 'TEXT PRIMARY KEY',
                'name': 'TEXT NOT NULL',
                'created_by': 'TEXT',
                'access_users': 'TEXT',
                'system_type': 'TEXT DEFAULT "personal"'
            }
        }
        
        changes_made = False
        
        for table_name, expected_columns in expected_schema.items():
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
            
            if not cursor.fetchone():
                print(f"Tabelle {table_name} in users.db existiert nicht - wird erstellt")
                if table_name == 'users':
                    cursor.execute('''CREATE TABLE IF NOT EXISTS users
                                     (id TEXT PRIMARY KEY, name TEXT NOT NULL)''')
                elif table_name == 'lager':
                    cursor.execute('''CREATE TABLE IF NOT EXISTS lager
                                     (id TEXT PRIMARY KEY, name TEXT NOT NULL, created_by TEXT,
                                       access_users TEXT, system_type TEXT DEFAULT 'personal')''')
                conn.commit()
                changes_made = True
                continue
            
            current_columns = get_table_columns(conn, table_name)
            
            for column_name, column_type in expected_columns.items():
                if column_name not in current_columns:
                    print(f"Fehlende Spalte in users.db erkannt: {table_name}.{column_name}")
                    
                    default_value = None
                    if 'DEFAULT' in column_type:
                        parts = column_type.split('DEFAULT')
                        column_type = parts[0].strip()
                        default_value = parts[1].strip().strip("'\"")
                        default_value = f"'{default_value}'"
                    
                    if add_missing_column(conn, table_name, column_name, column_type, default_value):
                        changes_made = True
        
        if changes_made:
            print("Migration für users.db abgeschlossen")
        else:
            print("Keine Änderungen für users.db erforderlich")
            
    except Exception as e:
        print(f"Fehler bei der users.db Migration: {e}")
    finally:
        conn.close()

def auto_migrate_all_databases():
    """Automatically check and migrate all databases"""
    print("=== Starte automatische Datenbank-Migration ===")
    
    # Migriere users.db
    check_and_migrate_users_db()
    
    # Finde alle Lager-Datenbanken
    for file in os.listdir('.'):
        if file.endswith('.db') and file != 'users.db':
            lager_id = file[:-3]  # Entferne .db
            print(f"\nPrüfe Lager-Datenbank: {lager_id}")
            check_and_migrate_warehouse_db(lager_id)
    
    print("\n=== Migration abgeschlossen ===")

def migrate_warehouse_db(lager_id):
       check_and_migrate_warehouse_db(lager_id)

def generate_random_id(length=6):
    return ''.join(random.choices(string.digits, k=length))

def get_db_connection(lager_id):
    return sqlite3.connect(f'{lager_id}.db')

def generate_qr_code(data):
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    img_str = base64.b64encode(buffer.getvalue()).decode()
    return f"data:image/png;base64,{img_str}"

def get_lager_system_type(lager_id):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT system_type FROM lager WHERE id = ?", (lager_id,))
    result = c.fetchone()
    conn.close()
    return result[0] if result else 'personal'


def backup_db(lager_id, operation):
    os.makedirs('backups', exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    shutil.copy(f'{lager_id}.db', f'backups/{timestamp}_{operation}_{lager_id}.db')

@app.route('/')
def login():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return render_template('login.html', title="Login")

@app.route('/login', methods=['POST'])
def do_login():
    user_id = request.form['user_id']
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT name FROM users WHERE id = ?", (user_id,))
    user = c.fetchone()
    conn.close()
    if user:
        session['user_id'] = user_id
        session['user_name'] = user[0]
        return redirect(url_for('dashboard'))
    else:
        return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT id, name FROM lager WHERE created_by = ? OR access_users LIKE ?",
              (session['user_id'], f"%{session['user_id']}%"))
    lagers = c.fetchall()
    conn.close()
    return render_template('dashboard.html', title="Dashboard", lagers=lagers, update_available=UPDATE_AVAILABLE, latest_version=LATEST_VERSION, current_version=VERSION)

@app.route('/create_lager', methods=['GET', 'POST'])
def create_lager():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        name = request.form['name']
        access_users = request.form.getlist('access_users')
        system_type = request.form.get('system_type', 'personal')
        lager_id = generate_random_id(8)
        while os.path.exists(f'{lager_id}.db'):
            lager_id = generate_random_id(8)
        conn = sqlite3.connect('users.db')
        c = conn.cursor()
        c.execute("INSERT INTO lager VALUES (?, ?, ?, ?, ?)", 
                  (lager_id, name, session['user_id'], ','.join(access_users), system_type))
        conn.commit()
        conn.close()
        create_warehouse_db(lager_id)
        return redirect(url_for('warehouse', lager_id=lager_id))
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT id, name FROM users WHERE id != ?", (session['user_id'],))
    users = c.fetchall()
    conn.close()
    return render_template('create_lager.html', title="Neues Lager", users=users)

@app.route('/lager/<lager_id>')
def warehouse(lager_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if not os.path.exists(f'{lager_id}.db'):
        return redirect(url_for('dashboard'))

    migrate_warehouse_db(lager_id)

    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT id FROM lager WHERE id = ? AND (created_by = ? OR access_users LIKE ?)",
              (lager_id, session['user_id'], f"%{session['user_id']}%"))
    has_access = c.fetchone()
    conn.close()

    if not has_access:
        return redirect(url_for('dashboard'))

    session['current_lager'] = lager_id
    return render_template('warehouse.html', title="Lager", lager_id=lager_id)

@app.route('/devices')
def devices():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    search = request.args.get('search', '')
    status_filters = [f for f in request.args.getlist('status') if f]
    art_filters = [f for f in request.args.getlist('art') if f]
    klasse_filters = [f for f in request.args.getlist('klasse') if f]
    sort_by = request.args.get('sort_by', 'name')
    group_by = request.args.get('group_by', 'none')  
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    
    query = """SELECT DISTINCT g.* FROM geraete g
               LEFT JOIN ausleih_details ad ON g.id = ad.geraet_id
               LEFT JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id AND a.status = 'ausgeliehen'
               WHERE 1=1"""
    params = []
    
    if search:
        query += " AND (g.name LIKE ? OR g.barcode LIKE ? OR g.lagerplatz LIKE ? OR g.seriennummer LIKE ? OR g.modell LIKE ? OR g.instrumentenart LIKE ? OR a.mitarbeiter_name LIKE ? OR a.klasse LIKE ?)"
        params.extend([f'%{search}%'] * 8)
    
    if status_filters:
        status_conditions = []
        if "verfügbar" in status_filters:
            status_conditions.append("g.status = 'verfügbar'")
        if "ausgeliehen" in status_filters:
            status_conditions.append("g.status LIKE 'ausgeliehen%'")
        if status_conditions:
            query += f" AND ({' OR '.join(status_conditions)})"
    
    if art_filters:
        placeholders = ','.join('?' for _ in art_filters)
        query += f" AND g.instrumentenart IN ({placeholders})"
        params.extend(art_filters)
    
    if klasse_filters:
        placeholders = ','.join('?' for _ in klasse_filters)
        query += f" AND a.klasse IN ({placeholders})"
        params.extend(klasse_filters)
    

    if sort_by == 'instrumentenart':
        query += " ORDER BY g.instrumentenart, g.name"
    elif sort_by == 'lagerplatz':
        query += " ORDER BY g.lagerplatz, g.name"
    elif sort_by == 'status':
        query += " ORDER BY g.status, g.name"
    elif sort_by == 'model':
        query += " ORDER BY g.modell, g.name"
    else:
        query += " ORDER BY g.name"
    
    c.execute(query, params)
    devices_list = c.fetchall()
    
    grouped_devices = {}
    if group_by == 'model':
        for device in devices_list:
            model = device[7] or 'Unbekanntes Modell'
            if model not in grouped_devices:
                grouped_devices[model] = []
            grouped_devices[model].append(device)
    elif group_by == 'series':
        for device in devices_list:
            instrument = device[9] or 'Unbekanntes Instrument'
            if instrument not in grouped_devices:
                grouped_devices[instrument] = []
            grouped_devices[instrument].append(device)
    elif group_by == 'serial':
        for device in devices_list:
            serial = device[6] or ''  
            first_letter = serial[0].upper() if serial else 'Unbekannt'
            if first_letter not in grouped_devices:
                grouped_devices[first_letter] = []
            grouped_devices[first_letter].append(device)
    elif group_by == 'instrument':
        for device in devices_list:
            instrument = device[8] or 'Unbekanntes Instrument'  
            if instrument not in grouped_devices:
                grouped_devices[instrument] = []
            grouped_devices[instrument].append(device)
    elif group_by == 'status':
        for device in devices_list:
            status = 'Verfügbar' if device[4] == 'verfügbar' else 'Ausgeliehen'
            if status not in grouped_devices:
                grouped_devices[status] = []
            grouped_devices[status].append(device)
    else:
        grouped_devices['Alle Geräte'] = devices_list
    
    c.execute("SELECT DISTINCT instrumentenart FROM geraete WHERE instrumentenart IS NOT NULL")
    instrumentenarten = [row[0] for row in c.fetchall()]
    c.execute("SELECT DISTINCT klasse FROM ausleihen WHERE klasse IS NOT NULL")
    klassen = [row[0] for row in c.fetchall()]
    conn.close()
    
    return render_template('devices.html', title="Geräte", devices=devices_list, grouped_devices=grouped_devices,
                         search=search, status_filters=status_filters,
                         art_filters=art_filters, klasse_filters=klasse_filters,
                         sort_by=sort_by, group_by=group_by, instrumentenarten=instrumentenarten, klassen=klassen)

@app.route('/add_device', methods=['GET', 'POST'])
def add_device():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        name = request.form['name']
        lagerplatz = request.form['lagerplatz']
        beschreibung = request.form['beschreibung']
        seriennummer = request.form['seriennummer']
        modell = request.form['modell']
        instrumentenart = request.form['instrumentenart']
        inventarnummer = request.form['inventarnummer']
        kaufdatum = request.form['kaufdatum']
        preis = request.form.get('preis', 0)
        quantity = int(request.form.get('quantity', 1)) if request.form.get('quantity_enabled') == 'on' else 1
        hersteller = request.form['hersteller']

        conn = get_db_connection(session['current_lager'])
        c = conn.cursor()

        while True:
            barcode = generate_random_id(6)
            c.execute("SELECT id FROM geraete WHERE barcode = ?", (barcode,))
            if not c.fetchone():
                break

        backup_db(session['current_lager'], 'before_add_device')
        c.execute("INSERT INTO geraete (name, barcode, lagerplatz, beschreibung, seriennummer, modell, instrumentenart, inventarnummer, kaufdatum, preis, quantity, hersteller) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                  (name, barcode, lagerplatz, beschreibung, seriennummer, modell, instrumentenart, inventarnummer, kaufdatum, preis, quantity, hersteller))
        conn.commit()
        backup_db(session['current_lager'], 'after_add_device')
        conn.close()
        return redirect(url_for('devices'))
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    c.execute("SELECT DISTINCT instrumentenart FROM geraete WHERE instrumentenart IS NOT NULL")
    instrumentenarten = [row[0] for row in c.fetchall()]
    conn.close()
    
    return render_template('add_device.html', title="Gerät hinzufügen", instrumentenarten=instrumentenarten)

@app.route('/edit_device/<int:device_id>', methods=['GET', 'POST'])
def edit_device(device_id):
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    
    if request.method == 'POST':
        c.execute("SELECT * FROM geraete WHERE id = ?", (device_id,))
        device = c.fetchone()

        name = request.form['name']
        barcode = request.form['barcode']
        lagerplatz = request.form['lagerplatz']
        new_beschreibung = request.form['beschreibung']
        seriennummer = request.form['seriennummer']
        modell = request.form['modell']
        instrumentenart = request.form['instrumentenart']
        inventarnummer = request.form['inventarnummer']
        kaufdatum = request.form['kaufdatum']
        preis = request.form.get('preis', 0)
        quantity = int(request.form.get('quantity', 1))
        hersteller = request.form['hersteller']
        defekt = request.form.get('defekt', 'off') == 'on'

        c.execute("SELECT id FROM geraete WHERE barcode = ? AND id != ?", (barcode, device_id))
        if c.fetchone():
            conn.close()
            return redirect(url_for('edit_device', device_id=device_id))

        current_beschreibung = device[5] or ''
        if new_beschreibung != current_beschreibung:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            user_name = session.get('user_name', 'Unknown')
            log_entry = f"[{timestamp} - {user_name}] {new_beschreibung}"
            if current_beschreibung:
                beschreibung = log_entry + "\n\n" + current_beschreibung
            else:
                beschreibung = log_entry
        else:
            beschreibung = current_beschreibung

        backup_db(session['current_lager'], 'before_edit_device')
        if defekt:
            status = 'defekt'
        else:
            status = 'verfügbar'  # temporary
        c.execute("UPDATE geraete SET name = ?, barcode = ?, lagerplatz = ?, beschreibung = ?, seriennummer = ?, modell = ?, instrumentenart = ?, inventarnummer = ?, kaufdatum = ?, preis = ?, quantity = ?, hersteller = ?, status = ? WHERE id = ?",
                  (name, barcode, lagerplatz, beschreibung, seriennummer, modell, instrumentenart, inventarnummer, kaufdatum, preis, quantity, hersteller, status, device_id))
        conn.commit()
        backup_db(session['current_lager'], 'after_edit_device')
        if not defekt:
            update_device_status(session['current_lager'], device_id)
        conn.close()
        return redirect(url_for('devices'))
    
    c.execute("SELECT * FROM geraete WHERE id = ?", (device_id,))
    device = c.fetchone()
    c.execute("SELECT DISTINCT instrumentenart FROM geraete WHERE instrumentenart IS NOT NULL")
    instrumentenarten = [row[0] for row in c.fetchall()]
    conn.close()
    
    beschreibung = device[5] or ''
    current_description = ''
    if beschreibung:
        entries = beschreibung.split('\n\n')
        if entries:
            latest_entry = entries[0]
            if latest_entry.startswith('[') and ']' in latest_entry:
                bracket_end = latest_entry.find(']') + 1
                current_description = latest_entry[bracket_end:].strip()
    
    return render_template('edit_device.html', title="Gerät bearbeiten", device=device, instrumentenarten=instrumentenarten, current_description=current_description)

@app.route('/delete_device/<int:device_id>')
def delete_device(device_id):
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    backup_db(session['current_lager'], 'before_delete_device')
    c.execute("DELETE FROM geraete WHERE id = ?", (device_id,))
    conn.commit()
    backup_db(session['current_lager'], 'after_delete_device')
    conn.close()
    return redirect(url_for('devices'))


@app.route('/admin/regenerate_missing_slips')
def admin_regenerate_missing_slips():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    result = regenerate_missing_borrow_pdfs()
    flash(result['message'], 'success' if result['success'] else 'error')
    return redirect(url_for('borrow'))

@app.route('/admin/download_all_slips')
def admin_download_all_slips():
    return get_all_borrows()

@app.route('/borrow_success/<ausleih_id>')
def borrow_success(ausleih_id):
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    system_type = get_lager_system_type(session['current_lager'])
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    
    c.execute("""SELECT a.ausleih_id, a.mitarbeiter_id, a.mitarbeiter_name, a.zielort, 
                        a.datum, a.rueckgabe_qr, a.status, a.email, a.klasse
                 FROM ausleihen a
                 WHERE a.ausleih_id = ?""", (ausleih_id,))
    borrow_row = c.fetchone()
    
    if not borrow_row:
        conn.close()
        return redirect(url_for('borrow'))
    
    borrow = {
        'ausleih_id': borrow_row[0],
        'mitarbeiter_id': borrow_row[1],
        'mitarbeiter_name': borrow_row[2],
        'zielort': borrow_row[3],
        'datum': borrow_row[4],
        'rueckgabe_qr': borrow_row[5],
        'status': borrow_row[6],
        'email': borrow_row[7],
        'klasse': borrow_row[8]
    }
    
    c.execute("""SELECT g.name, g.barcode, g.id, SUM(ad2.quantity) as total_quantity
                 FROM ausleih_details ad
                 JOIN geraete g ON ad.geraet_id = g.id
                 LEFT JOIN ausleih_details ad2 ON ad2.geraet_id = g.id
                 LEFT JOIN ausleihen a2 ON ad2.ausleih_id = a2.ausleih_id AND a2.status = 'ausgeliehen'
                 WHERE ad.ausleih_id = ?
                 GROUP BY g.id, g.name, g.barcode""", (ausleih_id,))
    devices = c.fetchall()
    conn.close()
    
    return render_template('borrow_success.html', title="Ausleihe erfolgreich", 
                         borrow=borrow, devices=devices, system_type=system_type, 
                         generate_qr_code=generate_qr_code)


@app.route('/regenerate_borrow_pdfs', methods=['POST'])
def regenerate_borrow_pdfs_route():
    """Route zum Regenerieren aller fehlenden Ausleih-PDFs"""
    if 'current_lager' not in session:
        return jsonify({'success': False, 'message': 'Kein Lager ausgewählt'})

    try:
        result = regenerate_missing_borrow_pdfs()
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Fehler: {str(e)}'})


def get_all_borrows():
    """Generate a ZIP file containing all borrow PDFs from the folder"""
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))

    lager_id = session['current_lager']
    pdf_dir = ensure_borrow_pdfs_directory(lager_id)

    # Alle PDF-Dateien im Ordner finden
    pdf_files = list(pdf_dir.glob("ausleihe_*.pdf"))

    if not pdf_files:
        return send_file(
            io.BytesIO(b'Keine Ausleihe-PDFs gefunden.'),
            mimetype='text/plain',
            as_attachment=True,
            download_name='keine_pdfs.txt'
        )

    # ZIP-Datei erstellen
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for pdf_file in pdf_files:
            # Dateinamen bereinigen für bessere Lesbarkeit
            clean_name = pdf_file.name
            zip_file.write(pdf_file, clean_name)

    zip_buffer.seek(0)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f'alle_ausleihen_{lager_id}_{timestamp}.zip'
    )


def _generate_borrow_pdf(self, borrow_info, devices):
    """Helper function to generate a single borrow PDF"""
    def get_base_name(name):
        match = re.match(r'^(.*?)\s*\d*$', name.strip())
        if match:
            base = match.group(1).strip()
            return base if base else name
        return name
    
    grouped_devices = {}
    for device in devices:
        device_id, name, barcode, modell, preis, seriennummer = device
        base_name = get_base_name(name)
        
        if base_name not in grouped_devices:
            grouped_devices[base_name] = {
                'name': base_name,
                'items': [],
                'count': 0,
                'total_price': 0,
                'model': modell or '',
                'image_path': None
            }
        
        grouped_devices[base_name]['items'].append({
            'id': device_id,
            'full_name': name,
            'barcode': barcode,
            'price': preis or 0,
            'seriennummer': seriennummer or ''
        })
        grouped_devices[base_name]['count'] += 1
        grouped_devices[base_name]['total_price'] += (preis or 0)
        
        if not grouped_devices[base_name]['image_path']:
            for ext in ['.jpg', '.png', '.jpeg']:
                img_path = f'images/{device_id}{ext}'
                if os.path.exists(img_path):
                    grouped_devices[base_name]['image_path'] = img_path
                    break
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'],
                                 fontSize=24, textColor=colors.HexColor('#1a1a1a'),
                                 spaceAfter=20, alignment=TA_LEFT)
    header_style = ParagraphStyle('Header', parent=styles['Normal'],
                                  fontSize=10, textColor=colors.HexColor('#666666'),
                                  spaceAfter=10)
    small_style = ParagraphStyle('Small', parent=styles['Normal'],
                                 fontSize=7, textColor=colors.HexColor('#666666'))
    
    elements = []
    elements.append(Paragraph(f"Ausleihe-Übersicht", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    info_text = f"""
    <b>Ausleihe-ID:</b> {borrow_info[0]}<br/>
    <b>Ausgeliehen an:</b> {borrow_info[1]}<br/>
    <b>Datum:</b> {borrow_info[2]}<br/>
    """
    if borrow_info[3]:
        info_text += f"<b>E-Mail:</b> {borrow_info[3]}<br/>"
    if borrow_info[4]:
        info_text += f"<b>Klasse:</b> {borrow_info[4]}<br/>"
    
    elements.append(Paragraph(info_text, header_style))
    elements.append(Spacer(1, 1*cm))
    
    table_data = [['', 'Artikel', 'S/N & QR', 'Menge', 'Preis/St.', 'Gesamt']]
    
    total_sum = 0
    for group_name, group_data in sorted(grouped_devices.items()):
        row = []
        
        if group_data['image_path']:
            try:
                img = RLImage(group_data['image_path'], width=2*cm, height=2*cm)
                row.append(img)
            except:
                row.append('')
        else:
            row.append('')
        
        name_text = f"<b>{group_data['name']}</b>"
        if group_data['model']:
            name_text += f"<br/><font size=8 color='#666666'>{group_data['model']}</font>"
        row.append(Paragraph(name_text, styles['Normal']))
        
        sn_qr_data = []
        for item in group_data['items']:
            sn = item['seriennummer'] if item['seriennummer'] else 'N/A'
            
            try:
                qr = qrcode.QRCode(version=1, box_size=3, border=1)
                qr.add_data(item['barcode'])
                qr.make(fit=True)
                qr_img = qr.make_image(fill_color="black", back_color="white")
                
                qr_buffer = BytesIO()
                qr_img.save(qr_buffer, format='PNG')
                qr_buffer.seek(0)
                
                qr_rl = RLImage(qr_buffer, width=1*cm, height=1*cm)
                
                mini_table = Table([[Paragraph(f"<font size=7>S/N: {sn}</font>", small_style), qr_rl]], 
                                   colWidths=[2*cm, 1.2*cm])
                mini_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 1),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                ]))
                sn_qr_data.append(mini_table)
            except:
                sn_qr_data.append(Paragraph(f"<font size=7>S/N: {sn}</font>", small_style))
        
        sn_qr_container = Table([[item] for item in sn_qr_data], colWidths=[3.2*cm])
        sn_qr_container.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 2),
            ('RIGHTPADDING', (0, 0), (-1, -1), 2),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        row.append(sn_qr_container)
        row.append(str(group_data['count']))
        
        avg_price = group_data['total_price'] / group_data['count'] if group_data['count'] > 0 else 0
        row.append(f"{avg_price:.2f} €")
        row.append(f"{group_data['total_price']:.2f} €")
        
        table_data.append(row)
        total_sum += group_data['total_price']
    
    table_data.append(['', Paragraph('<b>Summe</b>', styles['Normal']), '', '', '', f"{total_sum:.2f} €"])
    
    table = Table(table_data, colWidths=[2.5*cm, 6*cm, 3.5*cm, 1.5*cm, 2.5*cm, 2.5*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#333333')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (3, 0), (3, -1), 'CENTER'),
        ('ALIGN', (4, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -2), 0.5, colors.HexColor('#e0e0e0')),
        ('LINEABOVE', (0, -1), (-1, -1), 2, colors.HexColor('#333333')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    
    return buffer


def ensure_borrow_pdfs_directory(lager_id):
    pdf_dir = Path(f'borrow_pdfs_{lager_id}')
    pdf_dir.mkdir(exist_ok=True)
    return pdf_dir

def generate_borrow_pdf(borrow_info, devices):
    def get_base_name(name):
        match = re.match(r'^(.*?)\s*\d*$', name.strip())
        if match:
            base = match.group(1).strip()
            return base if base else name
        return name
    
    grouped_devices = {}
    for device in devices:
        device_id, name, barcode, modell, preis, quantity = device
        base_name = get_base_name(name)
        
        if base_name not in grouped_devices:
            grouped_devices[base_name] = {
                'name': base_name,
                'items': [],
                'count': 0,
                'total_price': 0,
                'model': modell or '',
                'image_path': None
            }
        
        grouped_devices[base_name]['items'].append({
            'id': device_id,
            'full_name': name,
            'barcode': barcode,
            'price': preis or 0,
            'quantity': quantity
        })
        grouped_devices[base_name]['count'] += quantity
        grouped_devices[base_name]['total_price'] += (preis or 0) * quantity
        
        if not grouped_devices[base_name]['image_path']:
            for ext in ['.jpg', '.png', '.jpeg']:
                img_path = f'images/{device_id}{ext}'
                if os.path.exists(img_path):
                    grouped_devices[base_name]['image_path'] = img_path
                    break
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=20,
        alignment=TA_LEFT
    )
    
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        spaceAfter=10
    )
    
    elements = []
    
    elements.append(Paragraph(f"Ausleihe-Übersicht", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    info_text = f"""
    <b>Ausleihe-ID:</b> {borrow_info[0]}<br/>
    <b>Ausgeliehen an:</b> {borrow_info[1]}<br/>
    <b>Datum:</b> {borrow_info[2]}<br/>
    """
    if borrow_info[3]:
        info_text += f"<b>E-Mail:</b> {borrow_info[3]}<br/>"
    if borrow_info[4]:
        info_text += f"<b>Klasse:</b> {borrow_info[4]}<br/>"
    
    elements.append(Paragraph(info_text, header_style))
    elements.append(Spacer(1, 1*cm))
    
    table_data = [['', 'Artikel', 'Menge', 'Preis/St.', 'Gesamt']]
    
    total_sum = 0
    for group_name, group_data in sorted(grouped_devices.items()):
        row = []
        
        if group_data['image_path']:
            try:
                img = RLImage(group_data['image_path'], width=2*cm, height=2*cm)
                row.append(img)
            except:
                row.append('')
        else:
            row.append('')
        
        name_text = f"<b>{group_data['name']}</b>"
        if group_data['model']:
            name_text += f"<br/><font size=8 color='#666666'>{group_data['model']}</font>"
        row.append(Paragraph(name_text, styles['Normal']))
        
        row.append(str(group_data['count']))
        
        avg_price = group_data['total_price'] / group_data['count'] if group_data['count'] > 0 else 0
        row.append(f"{avg_price:.2f} €")
        
        row.append(f"{group_data['total_price']:.2f} €")
        
        table_data.append(row)
        total_sum += group_data['total_price']
    
    table_data.append(['', Paragraph('<b>Summe</b>', styles['Normal']), '', '', f"{total_sum:.2f} €"])
    
    table = Table(table_data, colWidths=[3*cm, 8*cm, 2*cm, 3*cm, 3*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#333333')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (2, 0), (2, -1), 'CENTER'),
        ('ALIGN', (3, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -2), 0.5, colors.HexColor('#e0e0e0')),
        ('LINEABOVE', (0, -1), (-1, -1), 2, colors.HexColor('#333333')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    
    return buffer

def regenerate_missing_borrow_pdfs():
    if 'current_lager' not in session:
        return {'success': False, 'message': 'Kein Lager ausgewählt'}
    
    lager_id = session['current_lager']
    pdf_dir = ensure_borrow_pdfs_directory(lager_id)
    
    conn = get_db_connection(lager_id)
    c = conn.cursor()
    
    # Get all borrow IDs
    c.execute("SELECT ausleih_id FROM ausleihen")
    borrow_ids = [row[0] for row in c.fetchall()]
    
    generated = 0
    for ausleih_id in borrow_ids:
        pdf_path = pdf_dir / f"ausleihe_{ausleih_id}.pdf"
        
        # Generate PDF
        c.execute("""SELECT a.ausleih_id, a.mitarbeiter_name, a.datum, a.email, a.klasse
                     FROM ausleihen a
                     WHERE a.ausleih_id = ?""", (ausleih_id,))
        borrow_info = c.fetchone()
        
        c.execute("""SELECT g.id, g.name, g.barcode, g.modell, g.preis, ad.quantity
                     FROM ausleih_details ad
                     JOIN geraete g ON ad.geraet_id = g.id
                     WHERE ad.ausleih_id = ?
                     ORDER BY g.name""", (ausleih_id,))
        devices = c.fetchall()
        
        buffer = generate_borrow_pdf(borrow_info, devices)
        
        with open(pdf_path, 'wb') as f:
            f.write(buffer.getvalue())
        
        generated += 1
    
    conn.close()
    
    return {'success': True, 'message': f'{generated} PDFs regeneriert'}

app._generate_borrow_pdf = lambda borrow_info, devices: _generate_borrow_pdf(None, borrow_info, devices)

@app.route('/borrow_pdf/<ausleih_id>')
def borrow_pdf(ausleih_id):
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    
    c.execute("""SELECT a.ausleih_id, a.mitarbeiter_name, a.datum, a.email, a.klasse
                 FROM ausleihen a
                 WHERE a.ausleih_id = ?""", (ausleih_id,))
    borrow_info = c.fetchone()
    
    if not borrow_info:
        conn.close()
        return redirect(url_for('dashboard'))
    
    c.execute("""SELECT g.id, g.name, g.barcode, g.modell, g.preis, ad.quantity
                 FROM ausleih_details ad
                 JOIN geraete g ON ad.geraet_id = g.id
                 WHERE ad.ausleih_id = ?
                 ORDER BY g.name""", (ausleih_id,))
    devices = c.fetchall()
    conn.close()
    
    def get_base_name(name):
        """Extract base name without trailing numbers"""
        match = re.match(r'^(.*?)\s*\d*$', name.strip())
        if match:
            base = match.group(1).strip()
            return base if base else name
        return name
    
    grouped_devices = {}
    for device in devices:
        device_id, name, barcode, modell, preis, quantity = device
        base_name = get_base_name(name)

        if base_name not in grouped_devices:
            grouped_devices[base_name] = {
                'name': base_name,
                'items': [],
                'count': 0,
                'total_price': 0,
                'model': modell or '',
                'image_path': None
            }

        grouped_devices[base_name]['items'].append({
            'id': device_id,
            'full_name': name,
            'barcode': barcode,
            'price': preis or 0,
            'quantity': quantity
        })
        grouped_devices[base_name]['count'] += quantity
        grouped_devices[base_name]['total_price'] += (preis or 0) * quantity

        if not grouped_devices[base_name]['image_path']:
            for ext in ['.jpg', '.png', '.jpeg']:
                img_path = f'images/{device_id}{ext}'
                if os.path.exists(img_path):
                    grouped_devices[base_name]['image_path'] = img_path
                    break
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=20,
        alignment=TA_LEFT
    )
    
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        spaceAfter=10
    )
    
    elements = []
    
    elements.append(Paragraph(f"Ausleihe-Übersicht", title_style))
    elements.append(Spacer(1, 0.5*cm))
    

    info_text = f"""
    <b>Ausleihe-ID:</b> {borrow_info[0]}<br/>
    <b>Ausgeliehen an:</b> {borrow_info[1]}<br/>
    <b>Datum:</b> {borrow_info[2]}<br/>
    """
    if borrow_info[3]:
        info_text += f"<b>E-Mail:</b> {borrow_info[3]}<br/>"
    if borrow_info[4]:
        info_text += f"<b>Klasse:</b> {borrow_info[4]}<br/>"
    
    elements.append(Paragraph(info_text, header_style))
    elements.append(Spacer(1, 1*cm))
    
    table_data = [['', 'Artikel', 'Menge', 'Preis/St.', 'Gesamt']]
    
    total_sum = 0
    for group_name, group_data in sorted(grouped_devices.items()):
        row = []
        
        if group_data['image_path']:
            try:
                img = RLImage(group_data['image_path'], width=2*cm, height=2*cm)
                row.append(img)
            except:
                row.append('')
        else:
            row.append('')
        
        name_text = f"<b>{group_data['name']}</b>"
        if group_data['model']:
            name_text += f"<br/><font size=8 color='#666666'>{group_data['model']}</font>"
        row.append(Paragraph(name_text, styles['Normal']))
        
        row.append(str(group_data['count']))
        
        avg_price = group_data['total_price'] / group_data['count'] if group_data['count'] > 0 else 0
        row.append(f"{avg_price:.2f} €")

        row.append(f"{group_data['total_price']:.2f} €")
        
        table_data.append(row)
        total_sum += group_data['total_price']

    table_data.append(['', Paragraph('<b>Summe</b>', styles['Normal']), '', '', f"{total_sum:.2f} €"])

    table = Table(table_data, colWidths=[3*cm, 8*cm, 2*cm, 3*cm, 3*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#333333')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (2, 0), (2, -1), 'CENTER'),
        ('ALIGN', (3, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -2), 0.5, colors.HexColor('#e0e0e0')),
        ('LINEABOVE', (0, -1), (-1, -1), 2, colors.HexColor('#333333')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    
    return send_file(buffer, mimetype='application/pdf', 
                    as_attachment=True, download_name=f'ausleihe_{ausleih_id}.pdf')

def update_device_status(lager_id, device_id, conn=None):
    """
    Update device status based on current borrows.
    If conn is provided, use it. Otherwise create a new connection.
    """
    should_close = False
    if conn is None:
        conn = get_db_connection(lager_id)
        should_close = True
    
    try:
        c = conn.cursor()

        # Get device info
        c.execute("SELECT name, quantity FROM geraete WHERE id = ?", (device_id,))
        device = c.fetchone()
        if not device:
            return

        device_name, max_quantity = device

        # Calculate total borrowed per borrower
        c.execute("""SELECT a.mitarbeiter_name, SUM(ad.quantity) as total_quantity
                     FROM ausleih_details ad
                     JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id
                     WHERE ad.geraet_id = ? AND a.status = 'ausgeliehen'
                     GROUP BY a.mitarbeiter_name
                     ORDER BY a.mitarbeiter_name""", (device_id,))
        borrowers = c.fetchall()

        # Update status
        if not borrowers:
            status = 'verfügbar'
        else:
            borrower_strings = [f"{name} ({qty})" for name, qty in borrowers]
            status = ", ".join(borrower_strings)

        c.execute("UPDATE geraete SET status = ? WHERE id = ?", (status, device_id))
        
        if should_close:
            conn.commit()
            
    finally:
        if should_close and conn:
            conn.close()

@app.route('/borrow', methods=['GET', 'POST'])
def borrow():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    system_type = get_lager_system_type(session['current_lager'])
    
    if request.method == 'POST':
        if 'add_device' in request.form:
            barcode_inputs = request.form['barcode'].strip().split('\n')

            if 'borrow_list' not in session:
                session['borrow_list'] = []

            conn = get_db_connection(session['current_lager'])
            c = conn.cursor()

            added_count = 0
            error_messages = []

            for barcode_input in barcode_inputs:
                barcode = barcode_input.strip()
                if not barcode:
                    continue

                c.execute("SELECT id, name, barcode, quantity FROM geraete WHERE barcode = ?", (barcode,))
                device = c.fetchone()

                if device:
                    device_id, name, barcode_val, max_quantity = device
                    # Check how many are already borrowed
                    c.execute("SELECT SUM(ad.quantity) FROM ausleih_details ad JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id WHERE ad.geraet_id = ? AND a.status = 'ausgeliehen'", (device_id,))
                    already_borrowed = c.fetchone()[0] or 0
                    available = max_quantity - already_borrowed

                    if available <= 0:
                        error_messages.append(f"Gerät '{name}' ist nicht mehr verfügbar (alle {max_quantity} Exemplare ausgeliehen).")
                        continue

                    existing = next((d for d in session['borrow_list'] if d['id'] == device_id), None)
                    if existing:
                        if existing['quantity'] >= available:
                            error_messages.append(f"Maximale Anzahl für '{name}' bereits ausgewählt ({available}).")
                        else:
                            existing['quantity'] += 1
                            session.modified = True
                            added_count += 1
                    else:
                        session['borrow_list'].append({
                            'id': device_id, 'name': name, 'barcode': barcode_val,
                            'quantity': 1,
                            'max_quantity': available
                        })
                        session.modified = True
                        added_count += 1
                else:
                    error_messages.append(f"Gerät mit Barcode '{barcode}' nicht gefunden.")

            conn.close()

            if added_count > 0:
                flash(f"{added_count} Gerät(e) erfolgreich zur Ausleihliste hinzugefügt.", "success")
            if error_messages:
                for msg in error_messages:
                    flash(msg, "error")
                
                
        elif 'complete_borrow' in request.form:
            if system_type == 'personal':
                borrower_name = session['user_name']
                borrower_id = session['user_id']
                email = klasse = None
            else:
                borrower_name = request.form['borrower_name']
                borrower_id = request.form.get('borrower_id', 'N/A')
                email = request.form.get('email')
                klasse = request.form.get('klasse')
            
            if session.get('borrow_list'):
                ausleih_id = generate_random_id(4)
                
                # BACKUP VOR dem Öffnen der Connection
                backup_db(session['current_lager'], 'before_borrow')
                
                conn = get_db_connection(session['current_lager'])
                c = conn.cursor()
                
                c.execute("INSERT INTO ausleihen (ausleih_id, mitarbeiter_id, mitarbeiter_name, zielort, datum, rueckgabe_qr, email, klasse) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                          (ausleih_id, borrower_id, borrower_name, 'N/A',
                           datetime.now().strftime('%Y-%m-%d %H:%M:%S'), ausleih_id, email, klasse))
                
                for device in session['borrow_list']:
                    c.execute("INSERT INTO ausleih_details (ausleih_id, geraet_id, geraet_barcode, quantity) VALUES (?, ?, ?, ?)",
                              (ausleih_id, device['id'], device['barcode'], device['quantity']))
                    update_device_status(session['current_lager'], device['id'], conn)
                
                conn.commit()
                conn.close()
                
                # BACKUP NACH dem Schließen der Connection
                backup_db(session['current_lager'], 'after_borrow')
                
                session['borrow_list'] = []
                session.modified = True
                return redirect(url_for('borrow_success', ausleih_id=ausleih_id))
    
    borrow_list = session.get('borrow_list', [])
    return render_template('borrow.html', title="Ausleihen", borrow_list=borrow_list, system_type=system_type)

@app.route('/return', methods=['GET', 'POST'])
def return_devices():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    system_type = get_lager_system_type(session['current_lager'])
    
    if request.method == 'POST':
        if 'scan_qr' in request.form:
            qr_code = request.form['qr_code']
            return redirect(url_for('return_devices', qr=qr_code))
        elif 'complete_return' in request.form:
            ausleih_id = request.form['ausleih_id']
            device_ids = request.form.getlist('return_devices')
            
            # BACKUP VOR dem Öffnen der Connection
            backup_db(session['current_lager'], 'before_return')
            
            conn = get_db_connection(session['current_lager'])
            c = conn.cursor()
            
            for device_id in device_ids:
                c.execute("DELETE FROM ausleih_details WHERE ausleih_id = ? AND geraet_id = ?", (ausleih_id, device_id))
                update_device_status(session['current_lager'], device_id, conn)
            
            c.execute("SELECT COUNT(*) FROM ausleih_details WHERE ausleih_id = ?", (ausleih_id,))
            remaining = c.fetchone()[0]
            if remaining == 0:
                c.execute("UPDATE ausleihen SET status = 'zurückgegeben' WHERE ausleih_id = ?", (ausleih_id,))
            
            conn.commit()
            conn.close()
            
            # BACKUP NACH dem Schließen der Connection
            backup_db(session['current_lager'], 'after_return')
            
            return redirect(url_for('return_devices'))
    
    qr_code = request.args.get('qr')
    devices_to_return = []
    
    my_borrowed_devices = []
    if system_type == 'personal' and 'user_id' in session:
        conn = get_db_connection(session['current_lager'])
        c = conn.cursor()
        c.execute("""SELECT g.id, g.name, g.barcode, ad.ausleih_id, a.datum
                     FROM geraete g
                     JOIN ausleih_details ad ON g.id = ad.geraet_id
                     JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id
                     WHERE a.mitarbeiter_id = ? AND a.status = 'ausgeliehen'
                     ORDER BY a.datum DESC""", (session['user_id'],))
        my_borrowed_devices = c.fetchall()
        conn.close()
    
    if qr_code:
        conn = get_db_connection(session['current_lager'])
        c = conn.cursor()
        
        if system_type == 'personal':
            c.execute("""SELECT g.id, g.name, g.barcode, ad.ausleih_id
                         FROM geraete g
                         JOIN ausleih_details ad ON g.id = ad.geraet_id
                         JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id
                         WHERE a.rueckgabe_qr = ? AND a.status = 'ausgeliehen'""", (qr_code,))
        else:
            c.execute("""SELECT g.id, g.name, g.barcode, ad.ausleih_id
                         FROM geraete g
                         JOIN ausleih_details ad ON g.id = ad.geraet_id
                         JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id
                         WHERE g.barcode = ? AND a.status = 'ausgeliehen'""", (qr_code,))
        
        devices_to_return = c.fetchall()
        conn.close()
    
    return render_template('return.html', title="Zurückgeben", 
                         devices_to_return=devices_to_return, 
                         my_borrowed_devices=my_borrowed_devices,
                         system_type=system_type)

@app.route('/inventory')
def inventory():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    search = request.args.get('search', '')
    status_filters = [f for f in request.args.getlist('status') if f]
    art_filters = [f for f in request.args.getlist('art') if f]
    klasse_filters = [f for f in request.args.getlist('klasse') if f]
    sort_by = request.args.get('sort_by', 'name')
    group_by = request.args.get('group_by', 'none') 
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()

    query = """SELECT g.*, a.mitarbeiter_name, a.zielort, a.datum, a.email, a.klasse
               FROM geraete g
               LEFT JOIN ausleih_details ad ON g.id = ad.geraet_id
               LEFT JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id AND a.status = 'ausgeliehen'
               WHERE 1=1"""
    params = []
    
    if search:
        query += " AND (g.name LIKE ? OR g.barcode LIKE ? OR g.lagerplatz LIKE ? OR g.seriennummer LIKE ? OR g.modell LIKE ? OR g.instrumentenart LIKE ? OR a.mitarbeiter_name LIKE ? OR a.klasse LIKE ?)"
        params.extend([f'%{search}%'] * 8)
    
    if status_filters:
        status_conditions = []
        if "verfügbar" in status_filters:
            status_conditions.append("g.status = 'verfügbar'")
        if "ausgeliehen" in status_filters:
            status_conditions.append("g.status LIKE 'ausgeliehen%'")
        if status_conditions:
            query += f" AND ({' OR '.join(status_conditions)})"
    
    if art_filters:
        placeholders = ','.join('?' for _ in art_filters)
        query += f" AND g.instrumentenart IN ({placeholders})"
        params.extend(art_filters)
    
    if klasse_filters:
        placeholders = ','.join('?' for _ in klasse_filters)
        query += f" AND a.klasse IN ({placeholders})"
        params.extend(klasse_filters)

    if sort_by == 'instrumentenart':
        query += " ORDER BY g.instrumentenart, g.name"
    elif sort_by == 'lagerplatz':
        query += " ORDER BY g.lagerplatz, g.name"
    elif sort_by == 'status':
        query += " ORDER BY g.status, g.name"
    elif sort_by == 'model':
        query += " ORDER BY g.modell, g.name"
    else:
        query += " ORDER BY g.name"
    
    c.execute(query, params)
    devices_list = c.fetchall()
    
    grouped_devices = {}
    if group_by == 'model':
        for device in devices_list:
            model = device[7] or 'Unbekanntes Modell'
            if model not in grouped_devices:
                grouped_devices[model] = []
            grouped_devices[model].append(device)
    elif group_by == 'series':
        for device in devices_list:
            instrument = device[9] or 'Unbekanntes Instrument'
            if instrument not in grouped_devices:
                grouped_devices[instrument] = []
            grouped_devices[instrument].append(device)
    elif group_by == 'serial':
        for device in devices_list:
            serial = device[6] or ''
            first_letter = serial[0].upper() if serial else 'Unbekannt'
            if first_letter not in grouped_devices:
                grouped_devices[first_letter] = []
            grouped_devices[first_letter].append(device)
    elif group_by == 'instrument':
        for device in devices_list:
            instrument = device[8] or 'Unbekanntes Instrument'
            if instrument not in grouped_devices:
                grouped_devices[instrument] = []
            grouped_devices[instrument].append(device)
    elif group_by == 'status':
        for device in devices_list:
            status = 'Verfügbar' if device[4] == 'verfügbar' else 'Ausgeliehen'
            if status not in grouped_devices:
                grouped_devices[status] = []
            grouped_devices[status].append(device)
    else:
        grouped_devices['Alle Geräte'] = devices_list

    c.execute("SELECT DISTINCT instrumentenart FROM geraete WHERE instrumentenart IS NOT NULL")
    instrumentenarten = [row[0] for row in c.fetchall()]
    c.execute("SELECT DISTINCT klasse FROM ausleihen WHERE klasse IS NOT NULL")
    klassen = [row[0] for row in c.fetchall()]
    conn.close()
    
    return render_template('inventory.html', title="Inventar", devices=devices_list, grouped_devices=grouped_devices,
                         search=search, status_filters=status_filters,
                         art_filters=art_filters, klasse_filters=klasse_filters,
                         sort_by=sort_by, group_by=group_by, instrumentenarten=instrumentenarten, klassen=klassen)

@app.route('/remove_from_borrow/<int:device_id>')
def remove_from_borrow(device_id):
    if 'borrow_list' in session:
        session['borrow_list'] = [d for d in session['borrow_list'] if d['id'] != device_id]
        session.modified = True
    return redirect(url_for('borrow'))

@app.route('/manage_lager')
def manage_lager():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT id, name, created_by, access_users, system_type FROM lager WHERE created_by = ?", (session['user_id'],))
    lagers = c.fetchall()
    conn.close()
    
    return render_template('manage_lager.html', title="Lager verwalten", lagers=lagers)

@app.route('/edit_lager/<lager_id>', methods=['GET', 'POST'])
def edit_lager(lager_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    
    if request.method == 'POST':
        name = request.form['name']
        access_users = request.form.getlist('access_users')
        system_type = request.form.get('system_type', 'personal')
        c.execute("UPDATE lager SET name = ?, access_users = ?, system_type = ? WHERE id = ? AND created_by = ?",
                  (name, ','.join(access_users), system_type, lager_id, session['user_id']))
        conn.commit()
        conn.close()
        return redirect(url_for('manage_lager'))
    
    c.execute("SELECT * FROM lager WHERE id = ? AND created_by = ?", (lager_id, session['user_id']))
    lager = c.fetchone()
    if not lager:
        conn.close()
        return redirect(url_for('manage_lager'))
    
    c.execute("SELECT id, name FROM users WHERE id != ?", (session['user_id'],))
    users = c.fetchall()
    conn.close()
    
    return render_template('edit_lager.html', title="Lager bearbeiten", lager=lager, users=users)

@app.route('/delete_lager/<lager_id>')
def delete_lager(lager_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("DELETE FROM lager WHERE id = ? AND created_by = ?", (lager_id, session['user_id']))
    conn.commit()
    conn.close()
    
    if os.path.exists(f'{lager_id}.db'):
        os.remove(f'{lager_id}.db')
    
    return redirect(url_for('manage_lager'))

@app.route('/export')
def export():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    format_type = request.args.get('format', 'csv')
    
    if request.method == 'POST' or request.is_json:
        filter_data = request.get_json() if request.is_json else {}
        search = filter_data.get('search', '')
        status_filters = filter_data.get('status', [])
        art_filters = filter_data.get('art', [])
        klasse_filters = filter_data.get('klasse', [])
    else:
        search = request.args.get('search', '')
        status_filters = [f for f in request.args.get('status', '').split(',') if f]
        art_filters = [f for f in request.args.get('art', '').split(',') if f]
        klasse_filters = [f for f in request.args.get('klasse', '').split(',') if f]
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    
    query = """SELECT g.*, a.mitarbeiter_name, a.zielort, a.datum, a.email, a.klasse
               FROM geraete g
               LEFT JOIN ausleih_details ad ON g.id = ad.geraet_id
               LEFT JOIN ausleihen a ON ad.ausleih_id = a.ausleih_id AND a.status = 'ausgeliehen'
               WHERE 1=1"""
    params = []
    
    if search:
        query += " AND (g.name LIKE ? OR g.barcode LIKE ? OR g.lagerplatz LIKE ? OR g.seriennummer LIKE ? OR g.modell LIKE ? OR g.instrumentenart LIKE ? OR a.mitarbeiter_name LIKE ? OR a.klasse LIKE ?)"
        params.extend([f'%{search}%'] * 8)
    
    if status_filters:
        status_conditions = []
        if "verfügbar" in status_filters:
            status_conditions.append("g.status = 'verfügbar'")
        if "ausgeliehen" in status_filters:
            status_conditions.append("g.status LIKE 'ausgeliehen%'")
        if status_conditions:
            query += f" AND ({' OR '.join(status_conditions)})"
    
    if art_filters:
        placeholders = ','.join('?' for _ in art_filters)
        query += f" AND g.instrumentenart IN ({placeholders})"
        params.extend(art_filters)
    
    if klasse_filters:
        placeholders = ','.join('?' for _ in klasse_filters)
        query += f" AND a.klasse IN ({placeholders})"
        params.extend(klasse_filters)
    
    query += " ORDER BY g.instrumentenart, g.name"
    c.execute(query, params)
    devices_list = c.fetchall()
    conn.close()
    
    if format_type == 'csv':
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['Name', 'Barcode', 'Lagerplatz', 'Status', 'Beschreibung', 'Seriennummer', 'Modell', 'Hersteller', 'Instrumentenart', 'Inventar-Nummer', 'Kaufdatum', 'Preis', 'Ausgeliehen an', 'Email', 'Klasse'])

        for device in devices_list:
            writer.writerow([device[1], device[2], device[3], device[4], device[5], device[6], device[7], device[13] or '', device[8], device[9] or '', device[10] or '', device[11] or '', device[12] or '', device[15] or '', device[16] or ''])
        
        output.seek(0)
        return send_file(io.BytesIO(output.getvalue().encode('utf-8')), 
                        mimetype='text/csv', as_attachment=True, download_name='export.csv')
    
    elif format_type == 'word':
        doc = Document()
        doc.add_heading('Geräte Export', 0)
        
        table = doc.add_table(rows=1, cols=14)
        hdr_cells = table.rows[0].cells
        headers = ['Name', 'Barcode', 'Lagerplatz', 'Status', 'Beschreibung', 'Seriennummer', 'Modell', 'Instrumentenart', 'Inventar-Nummer', 'Kaufdatum', 'Preis', 'Ausgeliehen an', 'Email', 'Klasse']
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for device in devices_list:
            row_cells = table.add_row().cells
            row_cells[0].text = device[1]
            row_cells[1].text = device[2]
            row_cells[2].text = device[3]
            row_cells[3].text = device[4]
            row_cells[4].text = device[5] or ''
            row_cells[5].text = device[6] or ''
            row_cells[6].text = device[7] or ''
            row_cells[7].text = device[8] or ''
            row_cells[8].text = device[9] or ''
            row_cells[9].text = device[10] or ''
            row_cells[10].text = str(device[11] or '') + ' €'
            row_cells[11].text = device[12] or ''
            row_cells[12].text = device[15] or ''
            row_cells[13].text = device[16] or ''
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                        as_attachment=True, download_name='export.docx')
    
    elif format_type == 'pdf_labels':
        conn_layout = get_db_connection(session['current_lager'])
        c_layout = conn_layout.cursor()
        c_layout.execute("SELECT layout_data FROM label_layouts WHERE is_default = 1 LIMIT 1")
        layout_result = c_layout.fetchone()
        
        if not layout_result:
            c_layout.execute("SELECT layout_data FROM label_layouts ORDER BY created_at DESC LIMIT 1")
            layout_result = c_layout.fetchone()
        
        conn_layout.close()
        
        if not layout_result:
            return send_file(io.BytesIO(b'Kein Label-Layout gefunden. Bitte erstellen Sie zuerst ein Label-Layout.'), 
                        mimetype='text/plain', as_attachment=True, download_name='error.txt')
        
        layout_data = json.loads(layout_result[0])

        label_width_mm = float(layout_data.get('labelWidth', 50))
        label_height_mm = float(layout_data.get('labelHeight', 30))

        mm_to_pt = 2.834645669
        label_width_pt = label_width_mm * mm_to_pt
        label_height_pt = label_height_mm * mm_to_pt

        page_width, page_height = A4
        margin = 0.5 * cm

        usable_width = page_width - 2 * margin
        usable_height = page_height - 2 * margin

        labels_per_row = max(1, int(usable_width / label_width_pt))
        labels_per_col = max(1, int(usable_height / label_height_pt))
        labels_per_page = labels_per_row * labels_per_col
        
        buffer = BytesIO()
        
        def get_field_value(device, field_type):
            """Get the actual value for a field type from device data"""
            field_mapping = {
                'name': device[1],
                'barcode': device[2],
                'location': device[3],
                'status': device[4],
                'beschreibung': device[5] or '',
                'seriennummer': device[6] or '',
                'modell': device[7] or '',
                'instrumentenart': device[8] or '',
                'inventarnummer': device[9] or '',
                'kaufdatum': device[10] or '',
                'preis': f"{device[11] or 0} €",
                'borrower_name': device[12] or '',
                'destination': device[13] or '',
                'borrow_date': device[14] or '',
                'email': device[15] or '',
                'class': device[16] or '',
                'borrower_id': '',
                'text': ''
            }
            return field_mapping.get(field_type, '')
        
        def create_qr_code_image(data, width_pt, height_pt):
            """
            Generate QR code image and return as ReportLab Image object.
            
            Args:
                data: The data to encode in the QR code (e.g., barcode)
                width_pt: Width in points for the QR code image
                height_pt: Height in points for the QR code image
            
            Returns:
                ReportLab Image object ready to be drawn on canvas
            """
            import qrcode
            from reportlab.platypus import Image as RLImage
            from io import BytesIO
            
            qr = qrcode.QRCode(
                version=1, 
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=1,
            )
            
            qr.add_data(str(data))
            qr.make(fit=True)
            
            pil_img = qr.make_image(fill_color="black", back_color="white")
            
            img_buffer = BytesIO()
            pil_img.save(img_buffer, format='PNG')
            img_buffer.seek(0)

            return RLImage(img_buffer, width=width_pt, height=height_pt)
        
        def calculate_font_size(text, max_width_pt, max_height_pt, initial_font_size):
            """Calculate appropriate font size to fit text in given dimensions"""
            from reportlab.pdfbase.pdfmetrics import stringWidth
            
            font_size = initial_font_size
            min_font_size = 6
            
            while font_size >= min_font_size:
                text_width = stringWidth(str(text), 'Helvetica', font_size)
                
                if text_width <= (max_width_pt - 4):
                    return font_size
                
                font_size -= 1
            
            return min_font_size
        
        def create_single_label(device, canvas, x, y, width, height):
            """Draw a single label at the specified position"""
            
            canvas.setStrokeColor(colors.black)
            canvas.setLineWidth(0.5)
            canvas.rect(x, y, width, height)
            
            fields = layout_data.get('fields', [])
            
            px_to_pt = 0.75
            
            for field in fields:
                field_type = field.get('type', 'text')
                field_x_px = field.get('x', 0)
                field_y_px = field.get('y', 0)
                field_width_px = field.get('width', 100)
                field_height_px = field.get('height', 20)
                
                field_x_pt = x + (field_x_px * px_to_pt)
                field_y_pt = y + height - (field_y_px * px_to_pt) - (field_height_px * px_to_pt)
                field_width_pt = field_width_px * px_to_pt
                field_height_pt = field_height_px * px_to_pt
                
                if field_type == 'qr':
                    try:
                        qr_data = device[2]
                        qr_img = create_qr_code_image(qr_data, field_width_pt, field_height_pt)
                        qr_img.drawOn(canvas, field_x_pt, field_y_pt)
                    except Exception as e:
                        print(f"[v0] Error generating QR code: {e}")
                    
                else:
                    field_value = get_field_value(device, field_type)
                    
                    if field_type == 'text':
                        field_value = field.get('text', 'Text')
                    
                    if field_value:
                        font_size_raw = field.get('fontSize', '12px')
                        initial_font_size = 12
                        
                        if isinstance(font_size_raw, str):
                            match = re.search(r'(\d+)', font_size_raw)
                            if match:
                                initial_font_size = int(match.group(1))
                        elif isinstance(font_size_raw, (int, float)):
                            initial_font_size = int(font_size_raw)
                        
                        initial_font_size = max(6, min(initial_font_size, 24))
                        
                        font_size = calculate_font_size(str(field_value), field_width_pt, field_height_pt, initial_font_size)
                        
                        font_weight = field.get('fontWeight', 'normal')
                        font_name = 'Helvetica-Bold' if font_weight == 'bold' else 'Helvetica'
                        canvas.setFont(font_name, font_size)
                        

                        canvas.setFillColor(colors.black)
                        
                        text_align = field.get('textAlign', 'left')
                        
                        text_x = field_x_pt + 2 
                        if text_align == 'center':
                            text_x = field_x_pt + field_width_pt / 2
                        elif text_align == 'right':
                            text_x = field_x_pt + field_width_pt - 2
                        
                        text_y = field_y_pt + (field_height_pt / 2) - (font_size / 3)
                        
                        try:
                            if text_align == 'center':
                                canvas.drawCentredString(text_x, text_y, str(field_value))
                            elif text_align == 'right':
                                canvas.drawRightString(text_x, text_y, str(field_value))
                            else:
                                canvas.drawString(text_x, text_y, str(field_value))
                        except Exception as e:
                            print(f"[v0] Error drawing text: {e}")
        
        from reportlab.pdfgen import canvas as pdfcanvas
        
        pdf_canvas = pdfcanvas.Canvas(buffer, pagesize=A4)
        
        device_count = 0
        for device in devices_list:
            label_index = device_count % labels_per_page
            
            row = label_index // labels_per_row
            col = label_index % labels_per_row
            
            label_x = margin + (col * label_width_pt)
            label_y = page_height - margin - ((row + 1) * label_height_pt)
            
            create_single_label(device, pdf_canvas, label_x, label_y, label_width_pt, label_height_pt)
            
            device_count += 1

            if device_count % labels_per_page == 0 and device_count < len(devices_list):
                pdf_canvas.showPage()
        
        pdf_canvas.save()
        buffer.seek(0)
        
        return send_file(buffer, mimetype='application/pdf', 
                        as_attachment=True, download_name='etiketten.pdf')
        
    return send_file(io.BytesIO(b'Not implemented'), mimetype='text/plain')

@app.route('/label-layout')
def label_layout():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    c.execute("SELECT id, name, is_default, created_at FROM label_layouts ORDER BY is_default DESC, name")
    labels = c.fetchall()
    conn.close()
    
    return render_template('label_selection.html', title="Label Editor", labels=labels)

@app.route('/label-layout/edit/<int:label_id>')
def edit_label_layout(label_id):
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    c.execute("SELECT id, name, layout_data FROM label_layouts WHERE id = ?", (label_id,))
    label = c.fetchone()
    conn.close()
    
    if not label:
        return redirect(url_for('label_layout'))
    
    return render_template('label_layout.html', title=f"Label Editor - {label[1]}", 
                         label_id=label[0], label_name=label[1], layout_data=label[2])

@app.route('/label-layout/new')
def new_label_layout():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    return render_template('label_layout.html', title="Neues Label erstellen", 
                         label_id=None, label_name="", layout_data="{}")

@app.route('/save-layout', methods=['POST'])
def save_layout():
    if 'current_lager' not in session:
        return jsonify({'error': 'No warehouse selected'}), 400
    
    data = request.json
    label_id = data.get('label_id')
    label_name = data.get('name', 'Unbenannt')
    layout_data = json.dumps(data.get('layout', {}))
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    
    if label_id:
        c.execute("UPDATE label_layouts SET name = ?, layout_data = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                  (label_name, layout_data, label_id))
    else:
        c.execute("INSERT INTO label_layouts (name, layout_data) VALUES (?, ?)",
                  (label_name, layout_data))
        label_id = c.lastrowid
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'message': 'Layout gespeichert', 'label_id': label_id})

@app.route('/set-default-label/<int:label_id>', methods=['POST'])
def set_default_label(label_id):
    if 'current_lager' not in session:
        return jsonify({'error': 'No warehouse selected'}), 400
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    
    c.execute("UPDATE label_layouts SET is_default = 0")
    c.execute("UPDATE label_layouts SET is_default = 1 WHERE id = ?", (label_id,))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'message': 'Standard-Label gesetzt'})

@app.route('/delete-label/<int:label_id>', methods=['POST'])
def delete_label(label_id):
    if 'current_lager' not in session:
        return jsonify({'error': 'No warehouse selected'}), 400
    
    conn = get_db_connection(session['current_lager'])
    c = conn.cursor()
    c.execute("DELETE FROM label_layouts WHERE id = ?", (label_id,))
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'message': 'Label gelöscht'})

@app.route('/info', methods=['GET', 'POST'])
def info():
    if 'current_lager' not in session:
        return redirect(url_for('dashboard'))
    
    device_info = None
    borrow_info = None
    searched = False
    searched_code = None
    
    if request.method == 'POST' and 'search_device' in request.form:
        qr_code = request.form['qr_code'].strip()
        searched = True
        searched_code = qr_code
        
        if qr_code:
            conn = get_db_connection(session['current_lager'])
            c = conn.cursor()
              

            c.execute("SELECT * FROM geraete WHERE barcode = ?", (qr_code,))
            device = c.fetchone()
            
            if device:
                device_info = {
                    'id': device[0],
                    'name': device[1],
                    'barcode': device[2],
                    'lagerplatz': device[3],
                    'status': device[4],
                    'beschreibung': device[5],
                    'seriennummer': device[6],
                    'modell': device[7],
                    'instrumentenart': device[8],
                    'inventarnummer': device[9],
                    'kaufdatum': device[10],
                    'preis': device[11]
                }
                
                if device[4] != 'verfügbar' and device[4] != 'defekt':
                    c.execute("""SELECT a.* FROM ausleihen a
                                JOIN ausleih_details ad ON a.ausleih_id = ad.ausleih_id
                                WHERE ad.geraet_id = ? AND a.status = 'ausgeliehen'""", (device[0],))
                    borrow_data = c.fetchone()
                    
                    if borrow_data:
                        borrow_info = {
                            'ausleih_id': borrow_data[1],
                            'mitarbeiter_id': borrow_data[2],
                            'mitarbeiter_name': borrow_data[3],
                            'zielort': borrow_data[4],
                            'datum': borrow_data[5],
                            'rueckgabe_qr': borrow_data[6],
                            'email': borrow_data[8],
                            'klasse': borrow_data[9]
                        }
            
            conn.close()
    
    return render_template('info.html', title="Geräte-Info Scanner", 
                         device_info=device_info, borrow_info=borrow_info,
                         searched=searched, searched_code=searched_code)

@app.route('/update')
def update():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    try:
        import zipfile
        import tempfile
        import shutil

        # Vor jedem Update: Prüfe und lösche verschachteltes templates-Verzeichnis
        nested_templates = os.path.join('templates', 'templates')
        if os.path.exists(nested_templates):
            shutil.rmtree(nested_templates)

        response = requests.get("https://github.com/Matti-Krebelder/DMS/archive/refs/heads/main.zip")
        if response.status_code == 200:
            with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                tmp_file.write(response.content)
                tmp_file_path = tmp_file.name

            with zipfile.ZipFile(tmp_file_path, 'r') as zip_ref:
                app_py_info = zip_ref.getinfo('DMS-main/main.py')
                if app_py_info:
                    zip_ref.extract(app_py_info, '.')
                    if os.path.exists('DMS-main/main.py'):
                        shutil.move('DMS-main/main.py', 'main.py')

                for file_info in zip_ref.filelist:
                    if file_info.filename.startswith('DMS-main/templates/'):
                        relative_path = file_info.filename.replace('DMS-main/', '', 1)
                        if relative_path:
                            zip_ref.extract(file_info, '.')
                            extracted_path = file_info.filename
                            target_path = relative_path
                            if os.path.exists(extracted_path):
                                shutil.move(extracted_path, target_path)

            os.unlink(tmp_file_path)

            if os.path.exists('DMS-main'):
                shutil.rmtree('DMS-main')

            nested_templates = os.path.join('templates', 'templates')
            if os.path.exists(nested_templates):
                shutil.rmtree(nested_templates)

        else:
            return "Fehler beim Aktualisieren", 500

        return redirect(url_for('dashboard'))
    except Exception as e:
        return f"Fehler beim Update: {str(e)}", 500

@app.route('/logoutdms')
def logout():
    session.clear()
    return redirect(url_for('login'))

if __name__ == '__main__':
    init_user_db()
    check_version()
    app.run(debug=True, host='0.0.0.0', port=5000)
